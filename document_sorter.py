import docx
import os
import os.path
import shutil
from io import StringIO
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser

## ask for main directory
main_directory = input("\nPlease enter the path to the folder where your documents are located: ")
print("...")

## collect keywords until end token is received
END_TOKEN = '~'
keywords = []
print("\nEnter keywords one at a time. A folder will be created in the provided directory for each keyword and this is where documents with the given keyword will end up. \n\nWhen finished entering keywords, type {} and then press enter to proceed.\n".format(END_TOKEN))

while True:
    keyword = input("Type keyword and then press enter: ")
    keyword = keyword.lower()
    if keyword == END_TOKEN:
        print("Done, stopping keyword collection.")
        break

    keywords.append(keyword)

keyword_number = len(keywords)

## function for reversing list to iterate through to check for misses
def reverse(lst):
    return [ele for ele in reversed(lst)]

keywords_reversed = reverse(keywords)
print(keyword_number)

##intitialize document with docx library
document = docx.Document()

##function is for iterating through the .docx and append it to a string
def get_text(file_path):
    try:
        doc = docx.Document(file_path)
        fullText = []
        for para in doc.paragraphs:

            fullText.append(para.text.lower())

        return '\n'.join(fullText)
    except:
        pass

def convert_pdf_to_string(file_path):
    try:
        output_string = StringIO()
        with open(file_path, 'rb') as in_file:

            parser = PDFParser(in_file)
            doc = PDFDocument(parser)
            rsrcmgr = PDFResourceManager()
            device = TextConverter(rsrcmgr, output_string, laparams=LAParams())
            interpreter = PDFPageInterpreter(rsrcmgr, device)
            for page in PDFPage.create_pages(doc):

                interpreter.process_page(page)

        return(output_string.getvalue().lower())
    except:
        pass

#directory = r'/Users/nat/Library/Mobile Documents/com~apple~CloudDocs/Computer Science/Python/Document Sorter Test Files'
for keyword in keywords:
    path = os.path.join(main_directory, keyword)
    os.mkdir(path)

for filename in os.listdir(main_directory):

    if filename.endswith(".docx"):
        print("Checking {}".format(filename))
        global doc_fulltext
        doc_fulltext = get_text(os.path.join(main_directory, filename))
        for keyword in keywords:
            if doc_fulltext.find(keyword) > 0:
                try:
                    shutil.move(os.path.join(main_directory, filename), os.path.join(main_directory, keyword))
                    print("'{}' found in {}".format(keyword,filename) + ". Moving document now.")
                    try:
                        keywords_reversed.remove(keyword)
                    except ValueError:
                        continue
                except FileNotFoundError:
                    print("Cannot find file.")
            else:
                print("No keyword found in {}".format(filename))

    elif filename.endswith(".pdf"):
        print("Checking {}".format(filename))
        global pdf_fulltext
        pdf_fulltext = convert_pdf_to_string(os.path.join(main_directory, filename))
        for keyword in keywords:
            if pdf_fulltext.find(keyword) > 0:
                try:
                    shutil.move(os.path.join(main_directory, filename), os.path.join(main_directory, keyword))
                    print("'{}' found in {}".format(keyword,filename) + ". Moving document now.")
                    ## remove each keyword that is found from keywords_reversed
                    try:
                        keywords_reversed.remove(keyword)
                    except ValueError:
                        continue
                except FileNotFoundError:
                    print("Cannot find file.")
            else:
                print("No keyword found in {}".format(filename))
    else:
        continue

## iterating through files in reverse but in the created folders to check for more specific keywords that were missed in sequence
global i
i = 0
for filename in os.listdir(os.path.join(main_directory, keywords_reversed[i])):
    if filename.endswith(".docx"):
        print("Reverse checking {}".format(filename))
        doc_fulltext = get_text(os.path.join(main_directory, keyword, filename))
        for keyword in keywords_reversed:
            try:
                if doc_fulltext.find(keyword) > 0:
                    shutil.move(os.path.join(main_directory, keyword, filename), os.path.join(main_directory, keyword))
                    print("'{}' found in {}".format(keyword,filename) + ". Moving document now.")
                else:
                    print("No keyword found in {}".format(filename))
            except FileNotFoundError:
                print("Cannot find fike on reverse check.")
            except IndexError:
                print("Index out of range.")
            except AttributeError:
                pass



    elif filename.endswith(".pdf"):
        print("Reverse checking {}".format(filename))
        pdf_fulltext = convert_pdf_to_string(os.path.join(main_directory, filename))
        for keyword in keywords_reversed:
            try:
                if pdf_fulltext.find(keyword) > 0:
                    shutil.move(os.path.join(main_directory, keyword, filename), os.path.join(main_directory, keyword))
                    print("'{}' found in {}".format(keyword,filename) + ". Moving document now.")
                else:
                    print("No keyword found in {}".format(filename))
            except FileNotFoundError:
                print("Cannot find file on reverse check.")
            except AttributeError:
                pass
    else:
        continue
    i = i + 1
